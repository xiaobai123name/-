"""
多格式文档解析器
支持PDF、Word、Markdown格式的统一解析接口
"""

import os
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from abc import ABC, abstractmethod
import hashlib

# PDF解析
from pypdf import PdfReader
import pdfplumber

# Word解析
from docx import Document as DocxDocument

# Markdown解析
import markdown


@dataclass
class ParsedDocument:
    """解析后的文档数据结构"""
    content: str  # 文档全文内容
    pages: List[Dict[str, Any]]  # 按页/章节分割的内容
    metadata: Dict[str, Any]  # 文档元数据
    file_type: str  # 文件类型
    file_path: str  # 文件路径
    file_hash: str  # 文件内容哈希


class BaseParser(ABC):
    """解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档"""
        pass
    
    @staticmethod
    def compute_file_hash(file_path: str) -> str:
        """计算文件MD5哈希"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()


class PDFParser(BaseParser):
    """PDF文档解析器"""
    
    def __init__(self, use_llama_parse: bool = False, llama_api_key: Optional[str] = None):
        """
        初始化PDF解析器
        
        Args:
            use_llama_parse: 是否使用LlamaParse进行高级解析
            llama_api_key: LlamaParse API密钥
        """
        self.use_llama_parse = use_llama_parse
        self.llama_api_key = llama_api_key
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析PDF文档"""
        if self.use_llama_parse and self.llama_api_key:
            return self._parse_with_llama(file_path)
        else:
            return self._parse_with_pypdf(file_path)
    
    def _parse_with_pypdf(self, file_path: str) -> ParsedDocument:
        """使用PyPDF解析（基础方案）"""
        pages = []
        all_text = []
        
        # 使用pdfplumber获取更好的文本提取
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "content": text,
                    "tables": page.extract_tables() or []
                })
                all_text.append(text)
        
        # 获取元数据
        metadata = {}
        try:
            reader = PdfReader(file_path)
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "page_count": len(reader.pages)
                }
        except Exception:
            metadata = {"page_count": len(pages)}
        
        return ParsedDocument(
            content="\n\n".join(all_text),
            pages=pages,
            metadata=metadata,
            file_type="pdf",
            file_path=file_path,
            file_hash=self.compute_file_hash(file_path)
        )
    
    def _parse_with_llama(self, file_path: str) -> ParsedDocument:
        """使用LlamaParse解析（高级方案，支持表格和公式）"""
        try:
            from llama_parse import LlamaParse
            
            parser = LlamaParse(
                api_key=self.llama_api_key,
                result_type="markdown",  # 输出Markdown格式
                num_workers=1,
                verbose=False
            )
            
            # 解析文档
            documents = parser.load_data(file_path)
            
            # 合并所有内容
            all_text = "\n\n".join([doc.text for doc in documents])
            
            # 按页分割（LlamaParse可能返回整体内容）
            pages = []
            for i, doc in enumerate(documents):
                pages.append({
                    "page_number": i + 1,
                    "content": doc.text,
                    "metadata": doc.metadata if hasattr(doc, 'metadata') else {}
                })
            
            return ParsedDocument(
                content=all_text,
                pages=pages,
                metadata={"parser": "llama_parse", "page_count": len(pages)},
                file_type="pdf",
                file_path=file_path,
                file_hash=self.compute_file_hash(file_path)
            )
        except ImportError:
            # 如果llama_parse未安装，回退到基础方案
            return self._parse_with_pypdf(file_path)
        except Exception as e:
            # 解析失败，回退到基础方案
            print(f"LlamaParse解析失败，回退到PyPDF: {e}")
            return self._parse_with_pypdf(file_path)


class WordParser(BaseParser):
    """Word文档解析器"""
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文档"""
        doc = DocxDocument(file_path)
        
        pages = []
        all_text = []
        current_section = {"title": "", "content": [], "section_index": 0}
        section_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题样式
            if para.style.name.startswith('Heading'):
                # 保存上一个章节
                if current_section["content"]:
                    pages.append({
                        "section_index": current_section["section_index"],
                        "title": current_section["title"],
                        "content": "\n".join(current_section["content"])
                    })
                
                # 开始新章节
                section_index += 1
                current_section = {
                    "title": text,
                    "content": [],
                    "section_index": section_index
                }
            else:
                current_section["content"].append(text)
            
            all_text.append(text)
        
        # 保存最后一个章节
        if current_section["content"]:
            pages.append({
                "section_index": current_section["section_index"],
                "title": current_section["title"],
                "content": "\n".join(current_section["content"])
            })
        
        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # 元数据
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(tables),
            "section_count": len(pages),
            "tables": tables
        }
        
        # 尝试获取文档属性
        try:
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else ""
            })
        except Exception:
            pass
        
        return ParsedDocument(
            content="\n\n".join(all_text),
            pages=pages if pages else [{"section_index": 0, "title": "", "content": "\n".join(all_text)}],
            metadata=metadata,
            file_type="docx",
            file_path=file_path,
            file_hash=self.compute_file_hash(file_path)
        )


class MarkdownParser(BaseParser):
    """Markdown文档解析器"""
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析Markdown文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 按标题分割章节
        pages = []
        lines = content.split('\n')
        current_section = {"title": "", "content": [], "level": 0, "section_index": 0}
        section_index = 0
        
        for line in lines:
            # 检测Markdown标题
            if line.startswith('#'):
                # 计算标题级别
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                
                # 保存上一个章节
                if current_section["content"] or current_section["title"]:
                    pages.append({
                        "section_index": current_section["section_index"],
                        "title": current_section["title"],
                        "content": "\n".join(current_section["content"]),
                        "level": current_section["level"]
                    })
                
                # 开始新章节
                section_index += 1
                current_section = {
                    "title": title,
                    "content": [],
                    "level": level,
                    "section_index": section_index
                }
            else:
                current_section["content"].append(line)
        
        # 保存最后一个章节
        if current_section["content"] or current_section["title"]:
            pages.append({
                "section_index": current_section["section_index"],
                "title": current_section["title"],
                "content": "\n".join(current_section["content"]),
                "level": current_section["level"]
            })
        
        # 转换为HTML用于提取纯文本（可选）
        html_content = markdown.markdown(content)
        
        metadata = {
            "section_count": len(pages),
            "line_count": len(lines),
            "char_count": len(content)
        }
        
        return ParsedDocument(
            content=content,
            pages=pages if pages else [{"section_index": 0, "title": "", "content": content, "level": 0}],
            metadata=metadata,
            file_type="md",
            file_path=file_path,
            file_hash=self.compute_file_hash(file_path)
        )


class DocumentParser:
    """统一文档解析器接口"""
    
    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".md": "md",
        ".markdown": "md",
        ".txt": "txt"
    }
    
    def __init__(
        self, 
        use_llama_parse: bool = False, 
        llama_api_key: Optional[str] = None
    ):
        """
        初始化文档解析器
        
        Args:
            use_llama_parse: 是否使用LlamaParse解析PDF
            llama_api_key: LlamaParse API密钥
        """
        self.pdf_parser = PDFParser(use_llama_parse, llama_api_key)
        self.word_parser = WordParser()
        self.md_parser = MarkdownParser()
    
    def get_file_type(self, file_path: str) -> Optional[str]:
        """获取文件类型"""
        ext = Path(file_path).suffix.lower()
        return self.SUPPORTED_TYPES.get(ext)
    
    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持解析"""
        return self.get_file_type(file_path) is not None
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            ParsedDocument: 解析后的文档
            
        Raises:
            ValueError: 不支持的文件类型
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_type = self.get_file_type(file_path)
        if not file_type:
            raise ValueError(f"不支持的文件类型: {Path(file_path).suffix}")
        
        if file_type == "pdf":
            return self.pdf_parser.parse(file_path)
        elif file_type == "docx":
            return self.word_parser.parse(file_path)
        elif file_type in ("md", "txt"):
            return self.md_parser.parse(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
